import re 
import io
from pathlib import Path 
from typing import Optional 

import pdfplumber
from docx import Document 

from parsers.schemas import ParsedResume, Education, Experience, Project 

SECTION_KEYWORDS = {
    "skills":      ["skills", "technical skills", "technologies", "tech stack", "tools"],
    "experience":  ["experience", "work experience", "employment", "internships", "internship"],
    "education":   ["education", "academic background", "qualifications"],
    "projects":    ["projects", "personal projects", "academic projects", "key projects"],
    "summary":     ["summary", "objective", "profile", "about me", "about"],
    "certifications": ["certifications", "certificates", "courses", "achievements"],
    "languages":   ["languages", "spoken languages"],
}


SKILL_KEYWORDS = [
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust", "kotlin", "swift",
    "react", "next.js", "vue", "angular", "node.js", "express", "fastapi", "flask", "django",
    "mongodb", "postgresql", "mysql", "redis", "sqlite", "firebase",
    "docker", "kubernetes", "aws", "gcp", "azure", "jenkins", "github actions", "ci/cd",
    "git", "linux", "bash", "nginx", "gunicorn",
    "langchain", "langgraph", "huggingface", "pytorch", "tensorflow", "scikit-learn",
    "pandas", "numpy", "opencv", "rag", "llm", "openai", "claude",
    "playwright", "selenium", "rest api", "graphql", "websockets",
    "html", "css", "tailwind", "sass",
]


class ResumeParser:
    """
    Parses PDF or DOCX resumes into a structured ParsedResume object.
    Handles: contact info, skills, experience, education, projects, summary.
    """


    def parse(self, file_path: str | Path, file_bytes: Optional[bytes] = None) -> ParsedResume:
        """
        Entry point. Accepts either a file path or raw bytes + extension hint.
        Returns a fully populated ParsedResume dataclass.
        """
        path = Path(file_path) if file_path else None
        suffix = path.suffix.lower() if path else ""



        if file_bytes:
            # Detect type from magic bytes
            if file_bytes[:4] == b"%PDF":
                text = self._extract_text_pdf_bytes(file_bytes)
            else:
                text = self._extract_text_docx_bytes(file_bytes)
        elif suffix == ".pdf":
            text = self._extract_text_pdf(path)
        elif suffix in (".docx", ".doc"):
            text = self._extract_text_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}. Use PDF or DOCX.")
 
        return self._parse_text(text)
    



    def _extract_text_pdf(self, path: Path) -> str:
        with pdfplumber.open(str(path)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
 
    def _extract_text_pdf_bytes(self, data: bytes) -> str:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
 
    def _extract_text_docx(self, path: Path) -> str:
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
 
    def _extract_text_docx_bytes(self, data: bytes) -> str:
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    

    def _parse_text(self, raw_text: str) -> ParsedResume:
        lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
        sections = self._split_into_sections(lines)
 
        return ParsedResume(
            raw_text=raw_text,
            name=self._extract_name(lines),
            email=self._extract_email(raw_text),
            phone=self._extract_phone(raw_text),
            linkedin=self._extract_linkedin(raw_text),
            github=self._extract_github(raw_text),
            summary=self._extract_summary(sections),
            skills=self._extract_skills(raw_text, sections),
            experience=self._extract_experience(sections),
            education=self._extract_education(sections),
            projects=self._extract_projects(sections),
            certifications=self._extract_certifications(sections),
        )
    

    def _split_into_sections(self, lines: list[str]) -> dict[str, list[str]]:
        """
        Walk through lines and group them under their nearest section header.
        Returns dict like {"skills": [...], "experience": [...], ...}
        """
        sections: dict[str, list[str]] = {"header": []}
        current = "header"
 
        for line in lines:
            detected = self._detect_section(line)
            if detected:
                current = detected
                sections.setdefault(current, [])
            else:
                sections.setdefault(current, []).append(line)
 
        return sections
    

    def _detect_section(self, line: str) -> Optional[str]:
        clean = line.lower().strip(":-–—. ")
        for section, keywords in SECTION_KEYWORDS.items():
            if clean in keywords or any(clean == k for k in keywords):
                return section
        # Fuzzy: line is short + all caps → likely a header
        if len(line) < 35 and line.isupper():
            for section, keywords in SECTION_KEYWORDS.items():
                if any(k in line.lower() for k in keywords):
                    return section
        return None
    

    def _extract_name(self, lines: list[str]) -> str:
        """First non-empty line that looks like a name (no @ or digits, 2+ words)."""
        for line in lines[:6]:
            if "@" in line or re.search(r"\d{5,}", line):
                continue
            words = line.split()
            if 2 <= len(words) <= 5 and all(w[0].isupper() for w in words if w.isalpha()):
                return line
        return lines[0] if lines else "Unknown"
 
    def _extract_email(self, text: str) -> Optional[str]:
        match = re.search(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}", text)
        return match.group(0) if match else None
 
    def _extract_phone(self, text: str) -> Optional[str]:
        match = re.search(r"(\+?\d[\d\s\-().]{7,15}\d)", text)
        return match.group(0).strip() if match else None
 
    def _extract_linkedin(self, text: str) -> Optional[str]:
        match = re.search(r"linkedin\.com/in/[\w-]+", text, re.IGNORECASE)
        return f"https://{match.group(0)}" if match else None
 
    def _extract_github(self, text: str) -> Optional[str]:
        match = re.search(r"github\.com/[\w-]+", text, re.IGNORECASE)
        return f"https://{match.group(0)}" if match else None
 
    def _extract_summary(self, sections: dict) -> Optional[str]:
        lines = sections.get("summary", [])
        return " ".join(lines).strip() or None
 
    def _extract_skills(self, raw_text: str, sections: dict) -> list[str]:
        """
        Two-pass extraction:
        1. From the skills section lines (handles comma / pipe / bullet separated lists)
        2. Keyword scan across full text for anything missed
        """
        found: set[str] = set()
        skill_lines = sections.get("skills", [])
 
        for line in skill_lines:
            # Split on common delimiters
            parts = re.split(r"[,|•·/\n]", line)
            for p in parts:
                clean = p.strip().strip(":-").strip()
                if 1 < len(clean) < 40:
                    found.add(clean)
 
        # Keyword scan fallback
        text_lower = raw_text.lower()
        for kw in SKILL_KEYWORDS:
            if kw in text_lower and kw.title() not in found and kw not in found:
                found.add(kw)
 
        return sorted(found, key=str.lower)
 
    def _extract_experience(self, sections: dict) -> list[Experience]:
        lines = sections.get("experience", [])
        experiences: list[Experience] = []
        current: Optional[dict] = None
        bullets: list[str] = []
 
        for line in lines:
            # Detect a new role: line contains a year range or pipe-separated title|company
            if re.search(r"(19|20)\d{2}", line) or ("|" in line and len(line) < 80):
                if current:
                    experiences.append(Experience(**current, bullets=bullets))
                    bullets = []
                parts = re.split(r"[|–\-@]", line, maxsplit=3)
                current = {
                    "title": parts[0].strip() if len(parts) > 0 else line,
                    "company": parts[1].strip() if len(parts) > 1 else "",
                    "duration": parts[-1].strip() if len(parts) > 2 else "",
                }
            elif line.startswith(("•", "-", "*", "–")) or (current and len(line) > 20):
                bullets.append(line.lstrip("•-*– ").strip())
 
        if current:
            experiences.append(Experience(**current, bullets=bullets))
 
        return experiences
 
    def _extract_education(self, sections: dict) -> list[Education]:
        lines = sections.get("education", [])
        educations: list[Education] = []
        current: Optional[dict] = None
 
        for line in lines:
            # A new education entry: contains a degree keyword
            degree_keywords = ["b.e", "b.tech", "m.tech", "bsc", "msc", "bachelor", "master",
                               "b.e.", "be ", "mba", "phd", "diploma", "b.s", "m.s"]
            if any(kw in line.lower() for kw in degree_keywords):
                if current:
                    educations.append(Education(**current))
                year_match = re.search(r"(20\d{2})\s*[-–]\s*(20\d{2}|present)", line, re.IGNORECASE)
                current = {
                    "degree": line,
                    "institution": "",
                    "year": year_match.group(0) if year_match else "",
                    "gpa": "",
                }
            elif current and not current["institution"]:
                current["institution"] = line
            elif current:
                gpa_match = re.search(r"(gpa|cgpa|grade)[:\s]*([\d.]+)", line, re.IGNORECASE)
                if gpa_match:
                    current["gpa"] = gpa_match.group(2)
 
        if current:
            educations.append(Education(**current))
 
        return educations
 
    def _extract_projects(self, sections: dict) -> list[Project]:
        lines = sections.get("projects", [])
        projects: list[Project] = []
        current: Optional[dict] = None
        bullets: list[str] = []
        techs: list[str] = []
 
        for line in lines:
            # New project: short line that doesn't start with bullet
            if not line.startswith(("•", "-", "*")) and len(line) < 60 and len(line) > 3:
                if current:
                    projects.append(Project(**current, bullets=bullets, tech_stack=techs))
                    bullets, techs = [], []
                current = {"name": line, "description": ""}
                # inline tech hint e.g. "ProjectName | React, Node"
                if "|" in line:
                    parts = line.split("|", 1)
                    current["name"] = parts[0].strip()
                    techs = [t.strip() for t in parts[1].split(",")]
            elif line.lower().startswith(("tech", "stack", "built with", "technologies")):
                techs = [t.strip() for t in re.split(r"[,|:]", line)[1:]]
            elif current:
                bullets.append(line.lstrip("•-*– ").strip())
 
        if current:
            projects.append(Project(**current, bullets=bullets, tech_stack=techs))
 
        return projects
 
    def _extract_certifications(self, sections: dict) -> list[str]:
        lines = sections.get("certifications", [])
        return [l.lstrip("•-*– ").strip() for l in lines if len(l) > 3]
